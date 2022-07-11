# 엑셀로 저장된 수료 명단 정보를 불러와 워드로 수료증을 자동 생성하고 pdf로 변환
# pip install openpyxl
# pip install python-docx
# pip install docx2pdf
# ---------------------------------------------------------------------------------------------

# 파이썬 코드를 이용해 엑셀로 저장

# import pandas as pd # 데이터등을 다루는 라이브러리

# # 데이터 프레임 생성
# df = pd.DataFrame([["홍길동", "1990.01.02", "2021-0001"],
#                     ["김민준", "1990.05.06", "2021-0002"],
#                     ["김철수", "2000.08.08", "2021-0003"],
#                     ["김영희", "2000.09.09", "2021-0004"],
#                     ["이서준", "2010.10.10", "2021-0005"],
#                     ["장다인", "2017.12.12", "2021-0006"]])

# print(df)
# df.to_excel(r'.\수료증명단.xlsx', index=False, header=False) # 엑셀로 저장
# ---------------------------------------------------------------------------------------------

# 엑셀에 저장된 정보를 불러오는 프로그램

# from openpyxl import load_workbook # 엑셀파일을 읽어오기 위한 라이브러리
# load_wb = load_workbook(r'.\수료증명단.xlsx')
# load_ws = load_wb.active # 파일을 읽어온다 

# name_list = []
# birthday_list = []
# ho_list = []
# for i in range(1, load_ws.max_row +1): # 마지막 줄까지 for문 실행
#     name_list.append(load_ws.cell(i,1).value) # 1행 읽어 바인딩
#     birthday_list.append(load_ws.cell(i,2).value) # 2행 읽어 바인딩
#     ho_list.append(load_ws.cell(i,3).value) # 3행 읽어 바인딩
    
# print(name_list)
# print(birthday_list)
# print(ho_list)
# ---------------------------------------------------------------------------------------------

# 수료증 내용을 채운 후 저장하는 코드 작성
# 네이버에서 무료글꼴을 받아 저장해야 한다

import docx
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document(r'.\수료증양식.docx') # 양식 파일을 읽어옴

style = doc.styles['Normal'] # 기본이 되는 폰트와 글씨크기 정함
style.font.name = '나눔고딕'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
style.font.size = docx.shared.Pt(12)

para = doc.add_paragraph() # 문단생성 및 글씨입력 후 폰트와 글씨크기 정함
run = para.add_run('\n\n') 
run = para.add_run('              제 2020-0001 호\n') 
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20)

para = doc.add_paragraph()
run = para.add_run('\n\n')
run = para.add_run('수  료  증') 
run.font.name = '나눔고딕'
run.bold = True
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(40)
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

para = doc.add_paragraph()
run = para.add_run('\n\n')
run = para.add_run('        성       명: 장다인\n') 
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20)
run = para.add_run('        생 년 월 일: 2017.12.12\n') 
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20)
run = para.add_run('        교 육 과 정: 파이썬과 40개의 작품들\n')
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20) 
run = para.add_run('        교 육 날 짜: 2021.08.05~2021.09.09\n') 
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20)

para = doc.add_paragraph() 
run = para.add_run('\n\n')
run = para.add_run('        위 사람은 파이썬과 40개의 작품들 교육과정을\n') 
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20)
run = para.add_run('        이수하였으므로 이 증서를 수여 합니다.\n') 
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20)

para = doc.add_paragraph()
run = para.add_run('\n\n\n')
run = para.add_run('2021.09.19') 
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20)
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

para = doc.add_paragraph()
run = para.add_run('\n\n\n')
run = para.add_run('파이썬교육기관장') 
run.font.name = '나눔고딕'
run.bold = True
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20)
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save(r'.\수료증결과.docx')
# ---------------------------------------------------------------------------------------------

# 수료증 생성 후 pdf로 변환하는 코드 만들기

from docx import Document
from openpyxl import load_workbook
import docx
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert # pdf로 변환하기 위한 라이브러리

load_wb = load_workbook(r".\수료증명단.xlsx") # 여기부터
load_ws =load_wb.active

name_list = []
birthday_list = []
ho_list = []
for i in range(1,load_ws.max_row + 1):
    name_list.append(load_ws.cell(i, 1).value)
    birthday_list.append(load_ws.cell(i, 2).value)
    ho_list.append(load_ws.cell(i, 3).value)
    
print(name_list)
print(birthday_list)
print(ho_list) # 여기까지 엑셀의 값을 읽어오는코드

for i in range(len(name_list)): # 엑셀에서 읽은 이름 리스트의 길이만큼 반복
    doc = docx.Document(r'.\수료증양식.docx')
    style = doc.styles['Normal']
    style.font.name = '나눔고딕'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    style.font.size = docx.shared.Pt(12)
    
    para = doc.add_paragraph()
    run = para.add_run('\n\n') 
    run = para.add_run('              제 '+ ho_list[i] +' 호\n') # 수료증 맨 위의 호 저장
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)

    para = doc.add_paragraph()
    run = para.add_run('\n\n')
    run = para.add_run('수  료  증') 
    run.font.name = '나눔고딕'
    run.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(40)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para = doc.add_paragraph()
    run = para.add_run('\n\n')
    run = para.add_run('        성       명: ' + name_list[i] +'\n') # 수료증 성명 저장
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)
    run = para.add_run('        생 년 월 일: ' + birthday_list[i] +'\n') # 수료증 생년월일 저장
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)
    run = para.add_run('        교 육 과 정: 파이썬과 40개의 작품들\n')
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20) 
    run = para.add_run('        교 육 날 짜: 2021.08.05~2021.09.09\n') 
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)

    para = doc.add_paragraph()
    run = para.add_run('\n\n')
    run = para.add_run('        위 사람은 파이썬과 40개의 작품들 교육과정을\n') 
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)
    run = para.add_run('        이수하였으므로 이 증서를 수여 합니다.\n') 
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)

    para = doc.add_paragraph()
    run = para.add_run('\n\n\n')
    run = para.add_run('2021.09.19') 
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para = doc.add_paragraph()
    run = para.add_run('\n\n\n')
    run = para.add_run('파이썬교육기관장') 
    run.font.name = '나눔고딕'
    run.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save('.\\'+name_list[i]+'.docx') # 워드파일 생성
    convert('.\\'+name_list[i]+'.docx', # 워드파일 pdf로 저장
            '.\\'+name_list[i]+'.pdf')